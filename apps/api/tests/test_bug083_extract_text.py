"""BUG-083 — Subir minuta .docx daba 400 de Groq (file.text() sobre binario).

El front hacía `file.text()` sobre el .docx (un ZIP) → basura binaria que la
IA rechazaba con 400, reintentada 3 veces sin diagnóstico. Fixes:
- Extracción server-side con python-docx (`/api/v1/ai/extract-text`).
- `_call_ai_for_tenant` NO reintenta en 4xx (≠429) y propaga el body real.

Cubre:
- TC-083.1: extract_text_from_upload lee párrafos + tablas de un .docx.
- TC-083.2: texto plano se decodifica; .doc y formatos raros se rechazan.
- TC-083.3: endpoint /ai/extract-text devuelve el texto del .docx.
- TC-083.4: 4xx del provider NO se reintenta (1 sola llamada) y trae el body.
"""
from __future__ import annotations

import io

import httpx
import pytest

from app.services.document_text import (
    UnsupportedDocument,
    extract_text_from_upload,
)
from tests.factories import create_admin_role, create_tenant, create_user, login


def _build_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, val in enumerate(row):
                t.cell(r, c).text = val
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_tc083_1_extract_docx_paragraphs_and_tables():
    data = _build_docx(
        ["Acta de reunión semanal", "Asistentes: Ana, Beto"],
        table_rows=[["Acuerdo", "Responsable"], ["Cerrar sprint", "Ana"]],
    )
    text = extract_text_from_upload("minuta.docx", data)
    assert "Acta de reunión semanal" in text
    assert "Asistentes: Ana, Beto" in text
    # Las celdas de tabla (acuerdos/RAID) también se extraen.
    assert "Cerrar sprint" in text
    assert "Responsable" in text


def test_tc083_2_plain_and_unsupported():
    assert extract_text_from_upload("notas.txt", b"hola mundo") == "hola mundo"
    # sin extensión → se trata como texto plano.
    assert extract_text_from_upload("notas", b"plano") == "plano"
    # .doc legacy y formatos raros → error claro.
    with pytest.raises(UnsupportedDocument):
        extract_text_from_upload("viejo.doc", b"\xd0\xcf\x11\xe0")
    with pytest.raises(UnsupportedDocument):
        extract_text_from_upload("hoja.xlsx", b"PK\x03\x04")
    # un "docx" que no es un ZIP válido → error controlado, no crash.
    with pytest.raises(UnsupportedDocument):
        extract_text_from_upload("roto.docx", b"no soy un zip")


async def _admin(client, db_session):
    t = await create_tenant(db_session, slug="b83", name="B83")
    role = await create_admin_role(db_session, t)
    await create_user(
        db_session, tenant=t, username="b83admin",
        email="b83@acme.example.com", password="Str0ng-Admin-1!", roles=[role],
    )
    auth = await login(client, "b83admin", "Str0ng-Admin-1!")
    return t, auth


@pytest.mark.asyncio
async def test_tc083_3_extract_text_endpoint_docx(client, db_session):
    _t, auth = await _admin(client, db_session)
    data = _build_docx(["Reunión de kickoff", "Tema: alcance"])
    r = await client.post(
        "/api/v1/ai/extract-text",
        files={
            "file": (
                "kickoff.docx",
                io.BytesIO(data),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers=auth["_authz"],
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert "Reunión de kickoff" in body["text"]
    assert body["chars"] > 0


@pytest.mark.asyncio
async def test_tc083_3b_extract_text_endpoint_rejects_garbage_docx(client, db_session):
    _t, auth = await _admin(client, db_session)
    r = await client.post(
        "/api/v1/ai/extract-text",
        files={"file": ("roto.docx", io.BytesIO(b"no soy un zip"), "application/octet-stream")},
        headers=auth["_authz"],
    )
    assert r.status_code == 400, r.text
    assert r.json()["detail"]["code"] == "VALIDATION_ERROR"


@pytest.mark.asyncio
async def test_tc083_4_client_4xx_not_retried(monkeypatch):
    """Un 4xx del provider NO se reintenta (sería inútil) y el error final
    incluye el body real del 400 para diagnóstico."""
    from unittest.mock import AsyncMock, patch

    from app.workers.tasks import ai as ai_tasks
    from app.workers.tasks.ai import TenantAIConfig

    req = httpx.Request("POST", "https://api.groq.com/openai/v1/chat/completions")
    resp = httpx.Response(
        400,
        text='{"error":{"message":"context_length_exceeded","code":"context_length_exceeded"}}',
        request=req,
    )
    exc = httpx.HTTPStatusError("400 Bad Request", request=req, response=resp)

    gen = AsyncMock(side_effect=exc)
    with patch("app.workers.tasks.ai.generate_for_tenant", gen):
        with pytest.raises(RuntimeError) as ei:
            await ai_tasks._call_ai_for_tenant(
                "ping",
                system=None,
                tenant_cfg=TenantAIConfig(
                    mode="byo", byo={"provider": "openai", "api_key": "sk-x"}
                ),
                platform_groq_config=None,
                tenant_id="tenant-x",
                job_id="job-x",
            )
    # Una sola llamada: no se reintentó el 4xx.
    assert gen.await_count == 1
    # El error propaga la razón real del 400.
    assert "context_length_exceeded" in str(ei.value)
