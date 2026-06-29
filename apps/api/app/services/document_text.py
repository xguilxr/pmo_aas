"""BUG-083 — Extracción de texto plano de archivos subidos (minutas/transcripts).

El front sube un archivo (.docx / .txt / .srt / …) y necesita el texto para
mandarlo a la IA como `transcript`. Antes el front hacía `file.text()` sobre
el binario del .docx — que es un ZIP — produciendo basura binaria que Groq
rechazaba con `400 Bad Request`. Extraemos en el backend con `python-docx`
(ya es dependencia) para .docx, y `decode` para los formatos de texto plano.
"""
from __future__ import annotations

from io import BytesIO

# Formatos de texto plano que se decodifican directo.
PLAIN_TEXT_EXTS = {".txt", ".md", ".markdown", ".srt", ".vtt", ".csv", ".log", ".text"}
DOCX_EXT = ".docx"
SUPPORTED_EXTS = PLAIN_TEXT_EXTS | {DOCX_EXT}


class UnsupportedDocumentError(ValueError):
    """Extensión/contenido no soportado para extracción de texto."""


def file_ext(filename: str) -> str:
    """Extensión en minúsculas, incluido el punto (`'.docx'`); '' si no hay."""
    name = (filename or "").lower().strip()
    dot = name.rfind(".")
    return name[dot:] if dot >= 0 else ""


def extract_docx_text(data: bytes) -> str:
    """Texto de un .docx: párrafos del cuerpo + celdas de tablas.

    Los párrafos del cuerpo no incluyen el texto de las tablas, así que se
    recorren por separado (las minutas suelen tener acuerdos/RAID en tablas).
    """
    from docx import Document

    doc = Document(BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts).strip()


def _decode_plain(data: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc).strip()
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace").strip()


def extract_text_from_upload(filename: str, data: bytes) -> str:
    """Devuelve el texto plano de un archivo subido.

    Lanza `UnsupportedDocumentError` si el formato no se soporta o si el .docx no
    se puede leer (corrupto / no es un ZIP válido)."""
    ext = file_ext(filename)
    if ext == DOCX_EXT:
        try:
            return extract_docx_text(data)
        except UnsupportedDocumentError:
            raise
        except Exception as exc:  # docx corrupto / no es zip válido
            raise UnsupportedDocumentError(
                f"No se pudo leer el .docx ({type(exc).__name__}). "
                "Verifica que el archivo no esté dañado."
            ) from exc
    if ext in PLAIN_TEXT_EXTS or ext == "":
        return _decode_plain(data)
    if ext == ".doc":
        raise UnsupportedDocumentError(
            "El formato .doc (Word 97-2003) no se soporta. "
            "Guárdalo como .docx o pega el texto directamente."
        )
    raise UnsupportedDocumentError(
        f"Formato no soportado: {ext}. Usa .docx o texto plano "
        "(.txt, .md, .srt, .vtt), o pega el texto directamente."
    )
