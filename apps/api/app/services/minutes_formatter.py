"""Formateador estandarizado de Minuta (US-040, EP014, ENH-105).

ENH-105 — 6-section rigid structure (Highlander gold-standard):

    1. Encabezado
    2. Participantes (asistentes + ausentes justificados + no justificados)
    3. Resumen / Objetivo
    4. Temas tratados
    5. RAID unificado A/R/D/I
    6. Notas libres (opcional)

Las "actividades a hacer del backlog" se consolidan en RAID Acciones
(no hay sección separada — owner clarification 2026-05-22).

El formateador acepta tanto el shape nuevo (ENH-102 — 6 secciones con
`raid` como array A/R/D/I) como el legacy (agreements/risks/issues/
decisions con kinds). El nuevo se prefiere si la minuta ya guardó el
payload en ``raid_suggestions["_structured"]``.
"""
from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import date, datetime
from typing import Any

from app.core.tipografia import aplicar_a_docx
from app.models.modules import MeetingMinute
from app.models.project import Project
from app.services.ai.validator import validate_minute_payload
from app.services.pdf_renderer import render_pdf

SEPARATOR = "========"

# The 6 fixed sections in their canonical order — ENH-105 no reordering.
SECTION_ORDER: tuple[str, ...] = (
    "header",
    "participants",
    "summary",
    "topics",
    "raid",
    "free_notes",
)

RAID_TYPE_ORDER: tuple[str, ...] = ("A", "R", "D", "I")
RAID_TYPE_LABELS: dict[str, str] = {
    "A": "Acciones",
    "R": "Riesgos",
    "D": "Decisiones",
    "I": "Issues",
}


@dataclass
class MinuteView:
    """Vista canónica de la minuta — 6 secciones rígidas (ENH-105)."""

    title: str
    project_name: str | None
    project_folio: str | None
    header: dict[str, Any] = field(default_factory=dict)
    attendees: list[dict[str, Any]] = field(default_factory=list)
    absent_justified: list[dict[str, Any]] = field(default_factory=list)
    absent_unjustified: list[dict[str, Any]] = field(default_factory=list)
    summary: str = ""
    topics: list[dict[str, Any]] = field(default_factory=list)
    raid_by_type: dict[str, list[dict[str, Any]]] = field(default_factory=dict)
    free_notes: str = ""
    meeting_date: str = ""
    duration: str | None = None
    session_number: int | str | None = None

    def section_keys(self) -> tuple[str, ...]:
        return SECTION_ORDER

    def as_dict(self) -> dict[str, Any]:
        return {
            "title": self.title,
            "project_name": self.project_name,
            "project_folio": self.project_folio,
            "header": self.header,
            "attendees": self.attendees,
            "absent_justified": self.absent_justified,
            "absent_unjustified": self.absent_unjustified,
            "summary": self.summary,
            "topics": self.topics,
            "raid_by_type": self.raid_by_type,
            "raid_type_order": list(RAID_TYPE_ORDER),
            "raid_type_labels": RAID_TYPE_LABELS,
            "free_notes": self.free_notes,
            "meeting_date": self.meeting_date,
            "duration": self.duration,
            "session_number": self.session_number,
            "section_order": list(SECTION_ORDER),
        }


# ---- Format helpers ----------------------------------------------------------

def _format_meeting_date(value: Any) -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, datetime | date):
        return value.isoformat()
    return str(value)[:25]


def _short_title_from_minute(minute: MeetingMinute) -> str:
    raw = minute.title or "Reunión"
    if len(raw) > 60:
        raw = raw[:57].rstrip() + "…"
    return f'"Minuta Reunión {raw}"'


def _normalize_participant(p: Any) -> dict[str, Any]:
    if isinstance(p, dict):
        return {
            "name": p.get("name") or p.get("full_name") or "—",
            "role": p.get("role") or "",
            "area": p.get("area") or "",
            "email": p.get("email"),
        }
    return {"name": str(p), "role": "", "area": "", "email": None}


def _legacy_to_structured(minute: MeetingMinute) -> dict[str, Any]:
    """Build the 6-section payload from a legacy MeetingMinute row."""
    raid: list[dict[str, Any]] = []
    kind_to_type = {
        "action": "A", "acciones": "A", "agreement": "A",
        "risk": "R", "riesgo": "R",
        "issue": "I", "incident": "I", "incidente": "I",
        "decision": "D", "decisión": "D",
    }
    for ag in (minute.agreements or []):
        if not isinstance(ag, dict):
            ag = {"description": str(ag)}
        kind = (ag.get("kind") or ag.get("type") or "action").lower()
        t = kind_to_type.get(kind, "A")
        raid.append({
            "type": t,
            "description": ag.get("description") or ag.get("short_desc") or "—",
            "responsible": ag.get("owner_name") or ag.get("owner") or ag.get("suggested_owner_name"),
            "due_date": ag.get("due_date"),
            "status": ag.get("status") or "Open",
        })

    raid_block = minute.raid_suggestions or {}
    if isinstance(raid_block, dict):
        # Shape canónico A/R/D/I; legacy lessons/changes (descartados) no
        # se materializan en exports.
        for kind, code in (
            ("actions", "A"), ("risks", "R"),
            ("decisions", "D"), ("issues", "I"),
        ):
            for it in raid_block.get(kind, []) or []:
                if not isinstance(it, dict):
                    continue
                raid.append({
                    "type": code,
                    "description": it.get("short_desc") or it.get("description") or "—",
                    "responsible": it.get("suggested_owner_name") or it.get("owner_name"),
                    "due_date": it.get("suggested_due_date") or it.get("due_date"),
                    "status": it.get("status") if it.get("status") in {"Open", "In Progress", "Pending", "Closed"} else "Open",
                })

    topics: list[dict[str, Any]] = []
    for t in (minute.topics or []):
        if isinstance(t, dict):
            title = t.get("title") or t.get("name") or "—"
            bullets = t.get("bullets")
            if not isinstance(bullets, list):
                note = t.get("notes") or t.get("summary") or ""
                bullets = [note] if note else []
            topics.append({"title": title, "bullets": bullets})
        else:
            topics.append({"title": str(t), "bullets": []})

    return {
        "header": {},
        "participants": {
            "attendees": minute.participants or [],
            "absent_justified": [],
            "absent_unjustified": [],
        },
        "summary": "",
        "topics": topics,
        "raid": raid,
        "free_notes": None,
    }


def _extract_structured_payload(minute: MeetingMinute) -> dict[str, Any]:
    """Pick the 6-section payload from the minute.

    Order of precedence:
    1. ``minute.raid_suggestions["_structured"]`` if the ENH-102 pipeline
       persisted the structured 6-section payload there.
    2. Otherwise, synthesise from legacy fields.
    """
    raid_block = minute.raid_suggestions or {}
    raw_struct = None
    if isinstance(raid_block, dict) and isinstance(raid_block.get("_structured"), dict):
        raw_struct = raid_block["_structured"]
    if raw_struct is None:
        raw_struct = _legacy_to_structured(minute)
    normalized, _ = validate_minute_payload(raw_struct)
    return normalized


def build_view_from_payload(
    payload: dict[str, Any],
    *,
    title: str = "Minuta",
    project_name: str | None = None,
    project_folio: str | None = None,
    meeting_date: str = "",
) -> MinuteView:
    """Construye una MinuteView a partir de un payload 6-secciones."""
    normalized, _ = validate_minute_payload(payload)

    header = normalized.get("header") or {}
    participants = normalized.get("participants") or {}
    attendees = [_normalize_participant(p) for p in participants.get("attendees") or []]
    absent_j = [_normalize_participant(p) for p in participants.get("absent_justified") or []]
    absent_u = [_normalize_participant(p) for p in participants.get("absent_unjustified") or []]

    topics: list[dict[str, Any]] = []
    for idx, t in enumerate(normalized.get("topics") or [], start=1):
        if not isinstance(t, dict):
            continue
        bullets = t.get("bullets") or []
        if not isinstance(bullets, list):
            bullets = []
        topics.append({
            "index": idx,
            "title": t.get("title") or "—",
            "bullets": [str(b) for b in bullets if b],
        })

    raid_by_type: dict[str, list[dict[str, Any]]] = {k: [] for k in RAID_TYPE_ORDER}
    for item in normalized.get("raid") or []:
        raid_by_type[item["type"]].append(item)

    return MinuteView(
        title=title,
        project_name=project_name,
        project_folio=project_folio,
        header=header,
        attendees=attendees,
        absent_justified=absent_j,
        absent_unjustified=absent_u,
        summary=normalized.get("summary") or "",
        topics=topics,
        raid_by_type=raid_by_type,
        free_notes=normalized.get("free_notes") or "",
        meeting_date=meeting_date or (header.get("date") or ""),
        duration=header.get("duration"),
        session_number=header.get("session_number"),
    )


def build_view(minute: MeetingMinute, project: Project | None = None) -> MinuteView:
    """Transforma una MeetingMinute en la vista estandarizada (ENH-105)."""
    payload = _extract_structured_payload(minute)
    meeting_date_str = _format_meeting_date(minute.meeting_date)
    project_name = project.name if project else None
    project_folio = project.folio if project else None
    title_line = (
        f"{_short_title_from_minute(minute)} — "
        f"{project_name or 'Proyecto'} — {meeting_date_str[:10]}"
    )
    return build_view_from_payload(
        payload,
        title=title_line,
        project_name=project_name,
        project_folio=project_folio,
        meeting_date=meeting_date_str,
    )


# ---- Exports (6 secciones, mismo orden en TODOS los formatos) ---------------

def _section_label(idx: int, title: str) -> str:
    return f"{idx}. {title}"


def to_markdown(view: MinuteView) -> str:
    lines: list[str] = []
    lines.append(SEPARATOR)
    lines.append(f"Título — {view.title}")
    lines.append(SEPARATOR)

    # 1. Encabezado
    lines.append(_section_label(1, "Encabezado"))
    h = view.header or {}
    fields = [
        ("Fecha", h.get("date") or view.meeting_date),
        ("Hora", h.get("time")),
        ("Duración", h.get("duration") or view.duration),
        ("Tipo", h.get("type")),
        ("Facilitador", h.get("facilitator")),
        ("Modalidad", h.get("modality")),
        ("Sede", h.get("location")),
    ]
    for k, v in fields:
        if v:
            lines.append(f"  {k}: {v}")
    lines.append(SEPARATOR)

    # 2. Participantes
    lines.append(_section_label(2, "Participantes"))
    if view.attendees:
        lines.append("  Asistentes:")
        for p in view.attendees:
            role = f" — {p['role']}" if p.get("role") else ""
            area = f" — {p['area']}" if p.get("area") else ""
            lines.append(f"    - {p['name']}{role}{area}")
    if view.absent_justified:
        lines.append("  Ausentes justificados:")
        for p in view.absent_justified:
            lines.append(f"    - {p['name']}")
    if view.absent_unjustified:
        lines.append("  Ausentes no justificados:")
        for p in view.absent_unjustified:
            lines.append(f"    - {p['name']}")
    if not (view.attendees or view.absent_justified or view.absent_unjustified):
        lines.append("  —")
    lines.append(SEPARATOR)

    # 3. Resumen / Objetivo
    lines.append(_section_label(3, "Resumen / Objetivo"))
    lines.append(f"  {view.summary or '—'}")
    lines.append(SEPARATOR)

    # 4. Temas tratados
    lines.append(_section_label(4, "Temas tratados"))
    if view.topics:
        for t in view.topics:
            lines.append(f"  4.{t['index']} {t['title']}")
            for b in t["bullets"]:
                lines.append(f"    - {b}")
    else:
        lines.append("  —")
    lines.append(SEPARATOR)

    # 5. RAID A/R/D/I
    lines.append(_section_label(5, "RAID — A/R/D/I"))
    any_raid = False
    for tcode in RAID_TYPE_ORDER:
        items = view.raid_by_type.get(tcode) or []
        if not items:
            continue
        any_raid = True
        lines.append(f"  {tcode} — {RAID_TYPE_LABELS[tcode]}:")
        for it in items:
            resp = f" — {it['responsible']}" if it.get("responsible") else ""
            due = f" ({it['due_date']})" if it.get("due_date") else ""
            lines.append(f"    - {it['description']}{resp}{due} [{it.get('status', 'Open')}]")
    if not any_raid:
        lines.append("  —")
    lines.append(SEPARATOR)

    # 6. Notas libres
    lines.append(_section_label(6, "Notas libres"))
    lines.append(view.free_notes or "—")
    return "\n".join(lines) + "\n"


def to_plain_text(view: MinuteView) -> str:
    """Texto plano UTF-8 — alias estable de to_markdown."""
    return to_markdown(view)


def to_docx(view: MinuteView) -> bytes:
    from docx import Document

    doc = Document()
    # ENH-202: Helvetica en el estilo `Normal` del documento.
    aplicar_a_docx(doc)
    doc.add_heading(view.title, level=1)

    # 1. Encabezado
    doc.add_heading("1. Encabezado", level=2)
    h = view.header or {}
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid"
    rows = [
        ("Fecha", h.get("date") or view.meeting_date or "—"),
        ("Hora", h.get("time") or "—"),
        ("Duración", h.get("duration") or view.duration or "—"),
        ("Tipo", h.get("type") or "—"),
        ("Facilitador", h.get("facilitator") or "—"),
        ("Modalidad", h.get("modality") or "—"),
        ("Sede", h.get("location") or "—"),
    ]
    for k, v in rows:
        row = table.add_row()
        row.cells[0].text = k
        row.cells[1].text = str(v)

    # 2. Participantes
    doc.add_heading("2. Participantes", level=2)
    if view.attendees:
        doc.add_paragraph("Asistentes:", style="Intense Quote")
        for p in view.attendees:
            extra: list[str] = []
            if p.get("role"):
                extra.append(p["role"])
            if p.get("area"):
                extra.append(p["area"])
            suffix = f" — {' — '.join(extra)}" if extra else ""
            doc.add_paragraph(f"{p['name']}{suffix}", style="List Bullet")
    if view.absent_justified:
        doc.add_paragraph("Ausentes justificados:", style="Intense Quote")
        for p in view.absent_justified:
            doc.add_paragraph(p["name"], style="List Bullet")
    if view.absent_unjustified:
        doc.add_paragraph("Ausentes no justificados:", style="Intense Quote")
        for p in view.absent_unjustified:
            doc.add_paragraph(p["name"], style="List Bullet")
    if not (view.attendees or view.absent_justified or view.absent_unjustified):
        doc.add_paragraph("—")

    # 3. Resumen
    doc.add_heading("3. Resumen / Objetivo", level=2)
    doc.add_paragraph(view.summary or "—")

    # 4. Temas tratados
    doc.add_heading("4. Temas tratados", level=2)
    if view.topics:
        for t in view.topics:
            doc.add_heading(f"4.{t['index']} {t['title']}", level=3)
            for b in t["bullets"]:
                doc.add_paragraph(b, style="List Bullet")
    else:
        doc.add_paragraph("—")

    # 5. RAID
    doc.add_heading("5. RAID — A/R/D/I", level=2)
    any_raid = False
    for tcode in RAID_TYPE_ORDER:
        items = view.raid_by_type.get(tcode) or []
        if not items:
            continue
        any_raid = True
        doc.add_heading(f"{tcode} — {RAID_TYPE_LABELS[tcode]}", level=3)
        tbl = doc.add_table(rows=1 + len(items), cols=4)
        tbl.style = "Light Grid"
        headers = ["Descripción", "Responsable", "Fecha compromiso", "Status"]
        for j, c in enumerate(headers):
            tbl.rows[0].cells[j].text = c
        for i, it in enumerate(items, start=1):
            cells = tbl.rows[i].cells
            cells[0].text = it.get("description") or "—"
            cells[1].text = it.get("responsible") or "—"
            cells[2].text = str(it.get("due_date") or "—")
            cells[3].text = it.get("status") or "Open"
    if not any_raid:
        doc.add_paragraph("—")

    # 6. Notas libres
    doc.add_heading("6. Notas libres", level=2)
    doc.add_paragraph(view.free_notes or "—")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def to_pdf(view: MinuteView, tenant_name: str | None = None) -> bytes:
    context = view.as_dict()
    context["tenant_name"] = tenant_name
    return render_pdf("minutes/minute.html", context)
