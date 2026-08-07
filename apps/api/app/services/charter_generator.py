"""BUG-028 — Generador del Project Charter en .docx.

El owner (comentario 2026-04-24) pidió que cada proyecto tenga un
archivo de charter **real** (no un URL placeholder). Este servicio
genera el .docx desde la fila `ProjectCharter` usando `python-docx` y
lo persiste vía `document_storage.save_document_bytes()`, que usa el
backend configurado (`local` en dev, `s3` / Cloudflare R2 en prod
tras US-066).

El mismo archivo se edita desde dos lugares:
1. `/admin/projects/[id]/charter` (editor existente).
2. `/admin/projects/[id]/documents` (link al editor vía frontend).

Al guardar cambios en el editor, se llama a `regenerate_charter_docx()`
que sobrescribe el archivo (misma `Document.id`, bump de `version`).
"""
from __future__ import annotations

import logging
from datetime import UTC, datetime
from io import BytesIO
from uuid import UUID

import httpx
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.observabilidad import medido
from app.core.tipografia import aplicar_a_docx
from app.models.modules import Document
from app.models.organization import Organization
from app.models.project import Project
from app.models.project_charter import ProjectCharter
from app.services.branding_storage import find_logo_file
from app.services.document_storage import save_document_bytes
from app.services.folio import next_folio
from app.services.reports.scoped_status import HEALTH_COLOR

logger = logging.getLogger(__name__)

# python-docx 1.2.0 solo soporta add_picture con rasters; svg/webp rompen.
_DOCX_RASTER_EXTS = {"png", "jpg", "jpeg"}

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

# ENH-110: salud como color, no como palabra. En el .docx la representamos
# con un círculo "●" coloreado (font color RAG) en vez del texto.
# DAT-06 (2026-08-05): la clave `amber` salió. Era un alias del mismo color que
# `yellow`, y el glosario veta el término: la migración 0091 convirtió los datos
# a `yellow` a propósito, así que nada lo emite desde entonces. Un alias que
# nadie usa no es tolerancia, es el sitio por donde el vocabulario retirado
# vuelve la próxima vez que alguien copie este diccionario.
#
# DAT-05 (2026-08-05): y los colores eran los RETIRADOS. `#16a34a` y `#dc2626`
# están literalmente en la lista de retirados de `test_d7_paleta_de_salud.py`,
# que no los veía porque este archivo no figuraba entre los que pintan salud.
# O sea: D-7 unificó cuatro sitios y quedó un quinto, y el registro daba DAT-05
# por CONFORME. El acta de constitución en .docx —el documento que más se
# imprime y se firma— salía con la paleta anterior a DIS-02, la del verde que
# no llegaba a AA.
#
# Se derivan de `HEALTH_COLOR`, que es el origen, en vez de copiarse: una copia
# es lo que produjo las cinco.
_RAG_RGB = {
    estado: RGBColor.from_string(hex_color.lstrip("#").upper())
    for estado, hex_color in HEALTH_COLOR.items()
}


class _RagDot:
    """Marca un valor de tabla que debe renderizarse como ● de color."""

    def __init__(self, status: str | None) -> None:
        self.status = status


def _looks_raster(data: bytes) -> bool:
    """True si los bytes parecen PNG o JPEG (firma mágica)."""
    return data.startswith(b"\x89PNG\r\n\x1a\n") or data.startswith(b"\xff\xd8\xff")


def _load_local_tenant_logo(tenant_id: str) -> bytes | None:
    """ENH-111: bytes del logo del tenant desde disco (mismo origen que el
    branding y el email). Salta svg/webp (no soportados por python-docx)."""
    path = find_logo_file(tenant_id)
    if path is None:
        return None
    if path.suffix.lower().lstrip(".") not in _DOCX_RASTER_EXTS:
        return None
    try:
        return path.read_bytes()
    except OSError:
        return None


async def _download_image(url: str | None) -> bytes | None:
    """ENH-111: descarga una imagen raster desde una URL http(s) absoluta.
    Tolerante a fallos (timeout/red/formato): devuelve None y no rompe la
    generación del charter."""
    if not url or not url.lower().startswith(("http://", "https://")):
        return None
    try:
        async with httpx.AsyncClient(timeout=5.0, follow_redirects=True) as client:
            resp = await client.get(url)
        if resp.status_code != 200:
            return None
        ctype = resp.headers.get("content-type", "").lower()
        data = resp.content
        is_raster = ctype.startswith(("image/png", "image/jpeg", "image/jpg"))
        if not is_raster and not _looks_raster(data):
            return None
        return data
    except (httpx.HTTPError, OSError):
        return None


async def resolve_charter_logos(
    db: AsyncSession, tenant_id: UUID | str, project: Project
) -> list[bytes]:
    """ENH-111: resuelve los logos a insertar en el charter, en orden:
    tenant/PMO primero, cliente de la organización después. Cada logo puede
    venir como archivo local (tenant) o como URL externa (cliente, y tenant
    como fallback). Solo rasters; todo fallo degrada a 'sin ese logo'."""
    logos: list[bytes] = []
    tenant_id_s = str(tenant_id)

    tenant_logo = _load_local_tenant_logo(tenant_id_s)
    if tenant_logo is None:
        from app.models.tenant import Tenant

        tenant = (
            await db.execute(select(Tenant).where(Tenant.id == tenant_id_s))
        ).scalar_one_or_none()
        if tenant is not None:
            tenant_logo = await _download_image(tenant.logo_url)
    if tenant_logo:
        logos.append(tenant_logo)

    if project.organization_id:
        org = (
            await db.execute(
                select(Organization).where(
                    Organization.id == str(project.organization_id)
                )
            )
        ).scalar_one_or_none()
        if org is not None and org.client_logo_url:
            client_logo = await _download_image(org.client_logo_url)
            if client_logo:
                logos.append(client_logo)

    return logos


def _render_charter_docx(
    charter: ProjectCharter,
    project: Project,
    logos: list[bytes] | None = None,
) -> bytes:
    """Construye el .docx del charter usando python-docx y devuelve bytes.

    Estructura (secciones 1-3 de la tabla + sección 4 derivada del
    proyecto) que refleja la vista HTML en project_charters.py.

    `logos` (ENH-111): rasters tenant/cliente que se insertan centrados
    arriba del título. Resueltos por `resolve_charter_logos`.
    """
    doc = DocxDocument()
    # ENH-202: Helvetica en el estilo `Normal` del documento.
    aplicar_a_docx(doc)

    # ENH-111: portada con logo(s) arriba del título. Tolerante a imágenes
    # corruptas: si add_picture falla, se salta ese logo sin romper el docx.
    if logos:
        logo_par = doc.add_paragraph()
        logo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for idx, blob in enumerate(logos):
            try:
                if idx:
                    logo_par.add_run("    ")
                logo_par.add_run().add_picture(BytesIO(blob), width=Inches(1.4))
            except Exception:
                # Imagen inválida no debe romper la generación del charter.
                logger.warning("charter: logo %d inválido, se omite", idx)

    title = doc.add_heading(f"Project Charter — {charter.project_name}", level=1)
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0x18, 0x2E, 0x4E)

    subtitle = doc.add_paragraph()
    subrun = subtitle.add_run(
        f"Proyecto: {project.folio} · Generado {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    subrun.font.size = Pt(9)
    subrun.font.color.rgb = RGBColor(0x5A, 0x62, 0x77)

    def section(title_text: str, rows: list[tuple[str, object | None]]) -> None:
        doc.add_heading(title_text, level=2)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for label, value in rows:
            row = table.add_row()
            left = row.cells[0]
            right = row.cells[1]
            left.text = label
            for run in left.paragraphs[0].runs:
                run.bold = True
            if isinstance(value, _RagDot):
                # ENH-110: ● coloreado según el estado, sin texto.
                run = right.paragraphs[0].add_run("●")
                color = _RAG_RGB.get((value.status or "").lower())
                if color is not None:
                    run.font.color.rgb = color
            else:
                right.text = "" if value is None else str(value)

    section(
        "1. Información general",
        [
            ("Nombre", charter.project_name),
            ("Descripción", charter.description),
            ("Organización", charter.organization_id),
            ("Unidad de negocio", charter.business_unit_id),
            ("Departamento", charter.department_id),
        ],
    )

    section(
        "2. Stakeholders",
        [
            ("Sponsor", charter.sponsor),
            ("Email sponsor", charter.sponsor_email),
            ("Líder de negocio", charter.business_leader),
            ("Email líder negocio", charter.business_leader_email),
            ("Líder técnico", charter.tech_leader),
            ("Email líder técnico", charter.tech_leader_email),
            ("Project Manager", charter.pm_id),
        ],
    )

    section(
        "3. Clasificación y alcance",
        [
            ("Tipo", charter.project_type),
            ("Prioridad", charter.priority),
            ("Objetivo", charter.objective),
            ("Restricciones", charter.restrictions),
            ("Riesgos (resumen)", charter.risks_summary),
            ("Alcance", charter.scope),
            ("Personas clave", charter.key_people),
            ("Beneficios", charter.benefits),
        ],
    )

    section(
        "4. Gestión del proyecto",
        [
            ("Folio", project.folio),
            ("Fase", project.phase),
            ("Salud", _RagDot(project.health_status)),
            ("Inicio", project.start_date),
            ("Fin", project.end_date),
            ("Presupuesto", project.budget),
            ("Avance", f"{project.progress}%"),
        ],
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@medido("informe.acta_docx")
async def generate_charter_docx(
    db: AsyncSession,
    *,
    tenant_id: UUID | str,
    project: Project,
    charter: ProjectCharter,
    created_by: UUID | str | None = None,
) -> Document:
    """BUG-028 — crea/actualiza el `Document` category=charter con el
    archivo real en storage.

    Si ya existe un Document charter para el proyecto, **sobrescribe** el
    archivo (mismo Document.id) y bumpea `version`. No crea otro registro
    para evitar duplicados confusos en la vista de documentos.
    """
    tenant_id_s = str(tenant_id)
    existing = (
        await db.execute(
            select(Document).where(
                Document.project_id == str(project.id),
                Document.category == "charter",
                Document.is_current.is_(True),
                Document.deleted_at.is_(None),
            )
        )
    ).scalar_one_or_none()

    logos = await resolve_charter_logos(db, tenant_id_s, project)
    data = _render_charter_docx(charter, project, logos)

    if existing is None:
        folio = await next_folio(db, tenant_id=tenant_id_s, prefix="DOC")
        doc = Document(
            tenant_id=tenant_id_s,
            project_id=str(project.id),
            folio=folio,
            title=f"Project Charter — {project.name}",
            description="Documento fundacional del proyecto, generado automáticamente.",
            status="current",
            category="charter",
            mime_type=DOCX_CONTENT_TYPE,
            is_current=True,
            version=1,
            uploaded_by=str(created_by) if created_by else None,
            uploaded_at=datetime.now(UTC),
            created_by=str(created_by) if created_by else None,
        )
        db.add(doc)
        await db.flush()
    else:
        doc = existing
        doc.version = (doc.version or 1) + 1
        doc.mime_type = DOCX_CONTENT_TYPE
        doc.title = f"Project Charter — {project.name}"
        doc.uploaded_at = datetime.now(UTC)
        if created_by:
            doc.uploaded_by = str(created_by)

    file_url, _ = save_document_bytes(
        tenant_id_s,
        str(project.id),
        str(doc.id),
        data=data,
        ext="docx",
        content_type=DOCX_CONTENT_TYPE,
    )
    doc.file_url = file_url
    doc.size_bytes = len(data)
    await db.flush()
    return doc
